"""
Generate Disaster-ShieldAI Documentation
Creates a comprehensive Word document with all API endpoints and data sources
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

def create_documentation():
    """Generate the complete Disaster-ShieldAI documentation"""
    
    # Create document
    doc = Document()
    
    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_heading('Disaster-Shield AI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Complete System Documentation', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add version and date
    version_para = doc.add_paragraph('Version: 1.0.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. System Overview',
        '2. Architecture',
        '3. API Endpoints Reference',
        '   3.1 Risk Prediction Agent',
        '   3.2 Infrastructure Agent',
        '   3.3 Evacuation Agent',
        '   3.4 Citizen Intelligence Agent',
        '   3.5 Resource Allocation Agent',
        '   3.6 LangChain Orchestrator',
        '   3.7 Alerts (Twilio)',
        '   3.8 General Endpoints',
        '4. Data Sources Explained',
        '   4.1 Risk Prediction Agent Data',
        '   4.2 Infrastructure Agent Data',
        '   4.3 Evacuation Agent Data',
        '   4.4 Citizen Intelligence Agent Data',
        '   4.5 Resource Allocation Agent Data',
        '   4.6 LangChain Orchestrator Data',
        '   4.7 RAG Service Data',
        '5. System Workflow',
        '6. Deployment Guide',
        '7. Testing',
        '8. Troubleshooting'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()
    
    # ============================================================
    # 1. SYSTEM OVERVIEW
    # ============================================================
    doc.add_heading('1. System Overview', 1)
    
    doc.add_paragraph(
        'Disaster-Shield AI is a multi-agent AI-powered disaster management system '
        'for Sri Lanka. It integrates real-time weather, river gauge, and satellite '
        'data to predict flood and landslide risks, analyze infrastructure, plan '
        'evacuations, process citizen reports, and allocate emergency resources.'
    )
    
    # Key features
    doc.add_heading('1.1 Key Features', 2)
    features = [
        '• Real-time flood and landslide risk prediction (96.58% accuracy)',
        '• 335,419 road network analysis with ML-based status prediction',
        '• 100+ evacuation shelters with RAG-enabled search',
        '• Multilingual citizen intelligence (English, Sinhala, Tamil)',
        '• Resource optimization using Integer Linear Programming (ILP)',
        '• 25,123 historical disaster records from DesInventar (1974-2022)',
        '• RAG (Retrieval-Augmented Generation) with FAISS + Gemini Embeddings',
        '• LangChain orchestration with 6 tools',
        '• Twilio SMS and WhatsApp alerts',
        '• PostgreSQL database for structured data'
    ]
    for feature in features:
        doc.add_paragraph(feature)
    
    doc.add_page_break()
    
    # ============================================================
    # 2. ARCHITECTURE
    # ============================================================
    doc.add_heading('2. System Architecture', 1)
    
    doc.add_heading('2.1 Multi-Agent Architecture', 2)
    doc.add_paragraph(
        'The system consists of 5 specialized AI agents that work together '
        'through API calls and tool calling:'
    )
    
    agents = [
        ('Risk Prediction Agent', 'Predicts flood and landslide risk (0-100) using XGBoost + Random Forest'),
        ('Infrastructure Agent', 'Analyzes road conditions using ML models on 335,419 roads'),
        ('Evacuation Agent', 'Plans safe evacuation routes using A* pathfinding + RAG'),
        ('Citizen Intelligence Agent', 'Processes citizen reports using Gemini 3.5 (NLP, Vision, Translation)'),
        ('Resource Allocation Agent', 'Optimizes resource deployment using Integer Linear Programming')
    ]
    
    for name, desc in agents:
        doc.add_paragraph(f'• {name}: {desc}')
    
    doc.add_heading('2.2 Technology Stack', 2)
    tech_stack = [
        ('Backend', 'Python 3.10, Flask, FastAPI microservices'),
        ('Database', 'PostgreSQL'),
        ('AI/ML', 'XGBoost, Random Forest, Logistic Regression, scikit-learn'),
        ('LLM & AI APIs', 'Gemini 3.5, Gemini Embeddings, Google Maps Grounding'),
        ('Multi-Agent', 'LangChain (tool calling, RAG orchestration)'),
        ('Vector Store', 'FAISS'),
        ('Alerts', 'Twilio SMS, Twilio WhatsApp'),
        ('Frontend', 'React.js (planned)'),
        ('Hosting', 'GCP/AWS with Docker and GitHub Actions CI/CD')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    for tech, detail in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = detail
    
    doc.add_page_break()
    
    # ============================================================
    # 3. API ENDPOINTS REFERENCE
    # ============================================================
    doc.add_heading('3. API Endpoints Reference', 1)
    doc.add_paragraph(
        'All API endpoints are accessible at: http://localhost:5000'
    )
    
    # 3.1 Risk Prediction Agent
    doc.add_heading('3.1 Risk Prediction Agent', 2)
    doc.add_paragraph('Base URL: /api/risk')
    
    endpoints = [
        ('GET', '/districts', 'Get all 17 districts', '{"districts": [...], "count": 17}'),
        ('GET', '/predict/<district>', 'Predict risk for a specific district', '{"district": "Colombo", "prediction": {"risk_score": 73.53, "risk_level": "High"}}'),
        ('GET', '/predict-all', 'Predict risk for all districts', '{"total": 17, "predictions": [...]}'),
        ('GET', '/high-risk', 'Get high-risk districts (>70%)', '{"threshold": 70, "districts": [...]}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph('curl http://localhost:5000/api/risk/predict/Colombo', style='Intense Quote')
    
    # 3.2 Infrastructure Agent
    doc.add_heading('3.2 Infrastructure Agent', 2)
    doc.add_paragraph('Base URL: /api/infrastructure')
    
    endpoints = [
        ('GET', '/analyze/<district>', 'Analyze road conditions', '{"road_status": [...], "blocked_roads": 0}'),
        ('GET', '/road-status/<district>', 'Get road status summary', '{"total": 20, "roads": [...]}'),
        ('GET', '/status', 'Get agent status', '{"roads_monitored": 335419, "model_loaded": true}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph('curl http://localhost:5000/api/infrastructure/analyze/Colombo', style='Intense Quote')
    
    # 3.3 Evacuation Agent
    doc.add_heading('3.3 Evacuation Agent', 2)
    doc.add_paragraph('Base URL: /api/evacuation')
    
    endpoints = [
        ('GET', '/plan/<district>?lat=X&lon=Y', 'Plan evacuation route', '{"nearest_shelter": {...}, "evacuation_route": {...}}'),
        ('GET', '/shelters/<district>', 'Get shelters in district', '{"total_shelters": 12, "shelters": [...]}'),
        ('POST', '/route', 'Calculate custom route', '{"path": [...], "distance_km": 2.5}'),
        ('GET', '/status', 'Get agent status', '{"shelters_available": 100, "rag_enabled": true}'),
        ('POST', '/sync-shelters', 'Sync shelters to RAG', '{"synced_count": 34}'),
        ('POST', '/rag-search', 'RAG search shelters', '{"results": [...], "num_results": 5}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph('curl "http://localhost:5000/api/evacuation/plan/Colombo?lat=6.9271&lon=79.8612"', style='Intense Quote')
    
    # 3.4 Citizen Intelligence Agent
    doc.add_heading('3.4 Citizen Intelligence Agent', 2)
    doc.add_paragraph('Base URL: /api/citizen')
    
    endpoints = [
        ('POST', '/report/text', 'Submit text report', '{"report": {...}, "success": true}'),
        ('POST', '/report/image', 'Submit image report', '{"report": {...}}'),
        ('POST', '/report/voice', 'Submit voice report', '{"report": {...}}'),
        ('POST', '/chat', 'Chat with agent', '{"response": "...", "success": true}'),
        ('POST', '/translate', 'Translate text', '{"translated_text": "..."}'),
        ('GET', '/reports', 'Get all reports', '{"total": 5, "reports": [...]}'),
        ('GET', '/status', 'Get agent status', '{"gemini_enabled": true}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph(
        'curl -X POST http://localhost:5000/api/citizen/report/text \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"text": "Flooding in Colombo", "language": "en", "location": "Colombo"}\'',
        style='Intense Quote'
    )
    
    # 3.5 Resource Allocation Agent
    doc.add_heading('3.5 Resource Allocation Agent', 2)
    doc.add_paragraph('Base URL: /api/resource')
    
    endpoints = [
        ('POST', '/allocate', 'Allocate resources', '{"allocation_plan": [...], "resources_used": {...}}'),
        ('GET', '/status', 'Get resource status', '{"utilization": 42.66%}'),
        ('GET', '/deployment/<district>', 'Get deployment plan', '{"deployment_plan": {...}}'),
        ('GET', '/historical-summary', 'Get DesInventar summary', '{"total_events": 25123}'),
        ('POST', '/update', 'Update resources', '{"resources": {"rescue_teams": 5}}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph('curl -X POST http://localhost:5000/api/resource/allocate', style='Intense Quote')
    
    # 3.6 LangChain Orchestrator
    doc.add_heading('3.6 LangChain Orchestrator', 2)
    doc.add_paragraph('Base URL: /api/orchestrator')
    
    endpoints = [
        ('POST', '/chat', 'Multi-agent chat', '{"response": "...", "success": true}'),
        ('POST', '/chat-rag', 'RAG-enhanced chat', '{"response": "...", "rag_used": true}'),
        ('POST', '/knowledge/search', 'Search knowledge', '{"results": [...], "num_results": 5}'),
        ('GET', '/status', 'Get orchestrator status', '{"tools": 6, "llm_available": true}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph(
        'curl -X POST http://localhost:5000/api/orchestrator/chat \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"query": "What is the flood risk in Colombo?", "language": "en"}\'',
        style='Intense Quote'
    )
    
    # 3.7 Alerts
    doc.add_heading('3.7 Alerts (Twilio)', 2)
    doc.add_paragraph('Base URL: /api/alerts')
    
    endpoints = [
        ('POST', '/send', 'Send alert', '{"result": {...}}'),
        ('POST', '/send-bulk', 'Bulk alerts', '{"results": [...]}'),
        ('GET', '/status', 'Alert service status', '{"twilio_status": {...}}'),
        ('POST', '/test-sms', 'Test SMS', '{"success": true}'),
        ('POST', '/test-whatsapp', 'Test WhatsApp', '{"success": true}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response[:50] + '...'
    
    doc.add_paragraph()
    doc.add_paragraph('Example:')
    doc.add_paragraph(
        'curl -X POST http://localhost:5000/api/alerts/send \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"to": "+94787987255", "district": "Colombo", "risk_level": "High"}\'',
        style='Intense Quote'
    )
    
    # 3.8 General Endpoints
    doc.add_heading('3.8 General Endpoints', 2)
    
    endpoints = [
        ('GET', '/', 'System status', '{"name": "Disaster-ShieldAI", "version": "1.0.0"}'),
        ('GET', '/health', 'Health check', '{"status": "healthy"}'),
        ('GET', '/routes', 'All available routes', '{"routes": [...]}')
    ]
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Description'
    hdr[3].text = 'Sample Response'
    
    for method, endpoint, desc, response in endpoints:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = desc
        row[3].text = response
    
    doc.add_page_break()
    
    # ============================================================
    # 4. DATA SOURCES EXPLAINED
    # ============================================================
    doc.add_heading('4. Data Sources Explained', 1)
    
    # 4.1 Risk Prediction Agent
    doc.add_heading('4.1 Risk Prediction Agent Data', 2)
    
    data_sources = [
        ('Weather API', 'Rainfall, Temperature, Humidity, Wind Speed', 'OpenWeatherMap API (Real-time)'),
        ('River Gauges', 'Water levels in rivers', 'DMC CSV files (172 reports downloaded)'),
        ('Satellite', 'Flood extent, water levels', 'UNOSAT/Sentinel-1 Data'),
        ('Terrain Data', 'Elevation, Slope', 'SRTM Data'),
        ('Historical Data', 'Past flood events', 'DMC River gauges (917 records)')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How it works: Weather API + River Gauges + Satellite → Risk Model (XGBoost + Random Forest) → Risk Score (0-100)',
        style='Intense Quote'
    )
    
    # 4.2 Infrastructure Agent
    doc.add_heading('4.2 Infrastructure Agent Data', 2)
    
    data_sources = [
        ('Road Network', '335,419 roads', 'OpenStreetMap (Shapefile)'),
        ('ML Model', 'Road status prediction', 'Logistic Regression + Random Forest'),
        ('Risk Data', 'Risk scores', 'API call to Risk Agent')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How it works: Road Network + Risk Scores → ML Model → Road Status (Safe/Impassable/Blocked)',
        style='Intense Quote'
    )
    
    # 4.3 Evacuation Agent
    doc.add_heading('4.3 Evacuation Agent Data', 2)
    
    data_sources = [
        ('Shelters', '100+ shelters with lat/lon', 'CSV file (DMC data)'),
        ('Road Graph', '60,541 nodes, 51,480 edges', 'OpenStreetMap'),
        ('RAG', 'Shelter semantic search', 'FAISS + Gemini Embeddings'),
        ('Gemini', 'Real-time shelter search', 'Google Maps Grounding'),
        ('Pathfinding', 'A* algorithm', 'NetworkX graph')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How it works: Location → Find Nearest Shelter (RAG/Gemini) → A* Pathfinding → Safe Route',
        style='Intense Quote'
    )
    
    # 4.4 Citizen Intelligence Agent
    doc.add_heading('4.4 Citizen Intelligence Agent Data', 2)
    
    data_sources = [
        ('Gemini 3.5', 'NLP, Chat, Translation', 'Google Gemini API'),
        ('Knowledge Base', '4 disaster knowledge items', 'Gemini-generated + Fallback'),
        ('Feedback Loop', 'Citizen reports → Risk Agent', 'Internal API call')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How it works: Citizen Text/Image/Voice → Gemini NLP/Vision → Structured Report → Feedback to Risk Agent',
        style='Intense Quote'
    )
    
    # 4.5 Resource Allocation Agent
    doc.add_heading('4.5 Resource Allocation Agent Data', 2)
    
    data_sources = [
        ('DesInventar', '25,123 historical events (1974-2022)', 'CSV file'),
        ('Risk Data', 'Risk scores', 'API call to Risk Agent'),
        ('Resource Inventory', 'Available resources', 'CSV file')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph(
        'How it works: Historical Data + Risk Scores → ILP Optimization → Resource Allocation Plan',
        style='Intense Quote'
    )
    
    # 4.6 LangChain Orchestrator
    doc.add_heading('4.6 LangChain Orchestrator Data', 2)
    
    data_sources = [
        ('All Agents', '6 tools (all agents)', 'Internal API calls'),
        ('RAG', 'Knowledge retrieval', 'FAISS + Gemini Embeddings'),
        ('Gemini', 'LLM reasoning', 'Google Gemini API')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph(
        'How it works: User Query → Tool Calling → Multiple Agents → Coordinated Response',
        style='Intense Quote'
    )
    
    # 4.7 RAG Service
    doc.add_heading('4.7 RAG Service Data', 2)
    
    data_sources = [
        ('Knowledge Base', 'Disaster knowledge', 'JSON files + Gemini'),
        ('Shelter Data', 'Shelter information', 'CSV file'),
        ('FAISS', 'Vector search', 'FAISS index'),
        ('Gemini Embeddings', 'Text embeddings', 'Google Gemini API')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'What it provides'
    hdr[2].text = 'Where it comes from'
    
    for name, provides, source in data_sources:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = provides
        row[2].text = source
    
    doc.add_paragraph()
    doc.add_paragraph('Where RAG is used:')
    doc.add_paragraph('• Evacuation Agent: Find nearest shelters semantically', style='List Bullet')
    doc.add_paragraph('• Citizen Agent: Knowledge retrieval for chat', style='List Bullet')
    doc.add_paragraph('• Orchestrator: RAG-enhanced responses', style='List Bullet')
    doc.add_paragraph('• Knowledge Search: Search disaster knowledge base', style='List Bullet')
    
    doc.add_page_break()
    
    # ============================================================
    # 5. SYSTEM WORKFLOW
    # ============================================================
    doc.add_heading('5. System Workflow', 1)
    
    doc.add_heading('5.1 Complete Workflow Diagram', 2)
    doc.add_paragraph(
        'Citizen Report → Citizen Intelligence Agent → Feedback → Risk Prediction Agent → '
        'Risk Scores → Infrastructure Agent (Road Analysis) + Evacuation Agent (Shelter Search) → '
        'Resource Allocation Agent (Deployment) → Alerts (Twilio SMS/WhatsApp) → Dashboard'
    )
    
    doc.add_heading('5.2 Agent Communication Flow', 2)
    doc.add_paragraph(
        '1. Risk Agent predicts risk (0-100)\n'
        '2. If risk > 70%, Infrastructure Agent checks roads\n'
        '3. Evacuation Agent finds nearest shelter\n'
        '4. Resource Agent allocates resources\n'
        '5. Citizen Intelligence Agent processes reports\n'
        '6. Twilio sends alerts to citizens\n'
        '7. All data displayed on Dashboard'
    )
    
    doc.add_page_break()
    
    # ============================================================
    # 6. DEPLOYMENT GUIDE
    # ============================================================
    doc.add_heading('6. Deployment Guide', 1)
    
    doc.add_heading('6.1 Local Development', 2)
    doc.add_paragraph('1. Clone repository: git clone <repo-url>')
    doc.add_paragraph('2. Create virtual environment: python -m venv venv')
    doc.add_paragraph('3. Activate: source venv/bin/activate (Windows: venv\\Scripts\\activate)')
    doc.add_paragraph('4. Install dependencies: pip install -r requirements.txt')
    doc.add_paragraph('5. Setup PostgreSQL database')
    doc.add_paragraph('6. Configure .env file')
    doc.add_paragraph('7. Run server: python run.py')
    
    doc.add_heading('6.2 Environment Variables (.env)', 2)
    env_vars = [
        ('GEMINI_API_KEY', 'Your Google Gemini API key'),
        ('OPENWEATHER_API_KEY', 'Your OpenWeatherMap API key'),
        ('TWILIO_ACCOUNT_SID', 'Twilio Account SID'),
        ('TWILIO_AUTH_TOKEN', 'Twilio Auth Token'),
        ('TWILIO_PHONE_NUMBER', 'Twilio phone number (for SMS)'),
        ('DB_HOST', 'PostgreSQL host'),
        ('DB_NAME', 'PostgreSQL database name'),
        ('DB_USER', 'PostgreSQL username'),
        ('DB_PASSWORD', 'PostgreSQL password')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Variable'
    hdr[1].text = 'Description'
    
    for var, desc in env_vars:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = desc
    
    doc.add_page_break()
    
    # ============================================================
    # 7. TESTING
    # ============================================================
    doc.add_heading('7. Testing', 1)
    
    doc.add_heading('7.1 Run Complete Test Suite', 2)
    doc.add_paragraph('python scripts/test_complete_system.py')
    
    doc.add_heading('7.2 Test Results (23/23 passed)', 2)
    tests = [
        ('Risk Prediction Agent', '3/3', '✅'),
        ('Infrastructure Agent', '2/2', '✅'),
        ('Evacuation Agent', '3/3', '✅'),
        ('Citizen Intelligence Agent', '4/4', '✅'),
        ('Resource Allocation Agent', '3/3', '✅'),
        ('LangChain Orchestrator', '3/3', '✅'),
        ('Twilio Alerts', '3/3', '✅'),
        ('RAG Service', '2/2', '✅')
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Agent'
    hdr[1].text = 'Tests Passed'
    hdr[2].text = 'Status'
    
    for agent, passed, status in tests:
        row = table.add_row().cells
        row[0].text = agent
        row[1].text = passed
        row[2].text = status
    
    doc.add_page_break()
    
    # ============================================================
    # 8. TROUBLESHOOTING
    # ============================================================
    doc.add_heading('8. Troubleshooting', 1)
    
    doc.add_heading('8.1 Common Issues', 2)
    
    issues = [
        ('Server won\'t start', 'Check Python dependencies: pip install -r requirements.txt'),
        ('PostgreSQL connection error', 'Verify DB credentials in .env and ensure PostgreSQL is running'),
        ('Gemini API error', 'Check GEMINI_API_KEY in .env and ensure billing is enabled'),
        ('Twilio error', 'Verify TWILIO credentials in .env and phone number'),
        ('No roads loaded', 'Run: python scripts/download_roads.py'),
        ('No shelters found', 'Run: python scripts/download_evacuation_data.py'),
        ('RAG not working', 'Check: data/rag/vector_store/index.faiss exists')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Issue'
    hdr[1].text = 'Solution'
    
    for issue, solution in issues:
        row = table.add_row().cells
        row[0].text = issue
        row[1].text = solution
    
    doc.add_page_break()
    
    # ============================================================
    # APPENDIX
    # ============================================================
    doc.add_heading('Appendix: Key Metrics', 1)
    
    metrics = [
        ('Risk Prediction Accuracy', '96.58%'),
        ('Total Roads', '335,419'),
        ('Road Graph Nodes', '60,541'),
        ('Road Graph Edges', '51,480'),
        ('Shelters Available', '100+'),
        ('Historical Events', '25,123 (1974-2022)'),
        ('Districts Covered', '17'),
        ('Languages Supported', 'English, Sinhala, Tamil'),
        ('AI Models', 'XGBoost, Random Forest, Logistic Regression'),
        ('Tools in Orchestrator', '6'),
        ('Test Results', '23/23 passed')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'
    
    for metric, value in metrics:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value
    
    # ============================================================
    # SAVE THE DOCUMENT
    # ============================================================
    # Create output directory
    output_dir = 'documentation'
    os.makedirs(output_dir, exist_ok=True)
    
    # Save
    output_path = os.path.join(output_dir, 'DisasterShieldAI_Complete_Documentation.docx')
    doc.save(output_path)
    print(f"✅ Documentation saved to: {output_path}")
    print(f"📄 File size: {os.path.getsize(output_path) / 1024:.1f} KB")
    return output_path

if __name__ == "__main__":
    create_documentation()